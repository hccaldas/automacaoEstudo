import pyautogui as py
import time
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from bs4 import BeautifulSoup
import requests

    # Lista de sequencia da automação
# 1 - Abrir o chrome
# 2 - Maximizar janela
# 3 - Selecionar perfil do chrome - Trabalho
# 4 - Indo para facebook
# 5 - Checar se já esta logado
# 6 - Login facebook
# 7 - Aguardar carregamento da página e tentar encontrar a barra de buscas
# 8 - Buscar por Python
# 9 - Coletar posts
# 10 - Salvar posts

load_dotenv()


def salvar_imagem(nome_arquivo):
    if not os.path.exists('images2'):
        os.makedirs('images2')
    py.screenshot(f'images2/{nome_arquivo}')
    time.sleep(0.5)

def printar_post():
    # Criar pasta para as imagens se não existir
    if not os.path.exists('images2'):
        os.makedirs('images2')
        print("Pasta images2 criada")
    
    # Tirar prints da página
    print("\nTirando prints da página...")
    for i in range(10):
        print(f"Print {i+1}")
        
        # Salvar imagem atual
        py.screenshot(f'images2/page{i+1}.png')
        time.sleep(0.5)
        
        # Fazer scroll apenas se não for o último print
        if i < 10:
            py.scroll(-500)
            time.sleep(0.5)
    #volta para o inicio
    py.scroll(20000)
    time.sleep(0.5)

    # clica na barra de pesquisa do chrome
    py.click(1240, 48)
    time.sleep(0.5)
    py.hotkey('ctrl', 'c')
    time.sleep(0.5)

# Função para coletar posts usando PyAutoGUI
def coletar_post():
    print("\nIniciando coleta de posts usando PyAutoGUI...")
    posts_coletados = []
    
    try:
        # Buscar por Python
        print("Buscando por Python...")
        py.click(100, 100)  # Clicar em uma posição segura
        time.sleep(0.5)
        py.hotkey('ctrl', 'f')
        time.sleep(0.5)
        py.write("Python")
        time.sleep(0.5)
        py.press('enter')
        time.sleep(2)
        
        # Clicar em Posts
        print("Clicando em Posts...")
        py.click(42, 334)
        time.sleep(2)
        
        # Coletar posts
        print("Coletando posts...")
        
        # Fazer scroll para carregar mais posts
        for _ in range(5):  # Fazer 5 scrolls
            py.scroll(-500)
            time.sleep(2)
        
        # Coletar posts usando PyAutoGUI
        posts = list(py.locateAllOnScreen("images/post_content.png", confidence=0.7))
        
        if posts:
            print(f"Encontrados {len(posts)} posts!")
            
            for post in posts:
                try:
                    # Clicar no post
                    py.click(post)
                    time.sleep(1)
                    
                    # Copiar o conteúdo do post
                    py.hotkey('ctrl', 'c')
                    time.sleep(0.5)
                    
                    # Pegar o texto do clipboard
                    post_text = clipboard.paste()
                    
                    if post_text:
                        # Adicionar ao resultado
                        posts_coletados.append({
                            'conteudo': post_text
                        })
                        
                        print(f"Post coletado com sucesso!")
                    
                except Exception as e:
                    print(f"Erro ao processar post: {str(e)}")
                    continue
        
    except Exception as e:
        print(f"Erro ao coletar posts: {str(e)}")
    
    if not posts_coletados:
        print("\nNão foi possível coletar nenhum post")
    else:
        print(f"\nTotal de posts coletados: {len(posts_coletados)}")
        
        # Criar documento Word
        doc = Document()
        
        # Adicionar título
        titulo = doc.add_heading('Posts Coletados', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar cada post ao documento
        for i, post in enumerate(posts_coletados, 1):
            # Adicionar título do post
            doc.add_heading(f'Post {i}', level=1)
            
            # Adicionar conteúdo do post
            if post['conteudo']:
                paragrafo = doc.add_paragraph()
                run = paragrafo.add_run(post['conteudo'])
                run.font.size = Pt(12)
            else:
                paragrafo = doc.add_paragraph("Conteúdo não disponível")
            
            # Adicionar linha separadora
            doc.add_paragraph('------------------------')
        
        # Salvar documento
        doc.save('posts_coletados.docx')
        print("Posts salvos com sucesso no arquivo 'posts_coletados.docx'")
    
    return posts_coletados

def login_facebook():
    # Configuração
    py.PAUSE = 1  # Adiciona uma pausa de 1 segundo entre as ações
    py.FAILSAFE = True  # Permite mover o mouse para o canto superior esquerdo para interromper o script

    # Apertar win e digitar chrome
    py.hotkey('win', 'r')
    py.write('chrome')
    py.press('enter')
    time.sleep(2)

    # Selecionar perfil do chrome - Trabalho
    teclaTab = 'tab'
    numero_de_vezes = 10
    py.press(teclaTab, numero_de_vezes)
    time.sleep(0.5)
    py.press('enter')
    time.sleep(0.5)

    py.hotkey('win', 'up')
    time.sleep(0.5)

    #indo para facebook
    print("Indo para facebook")
    py.write("https://www.facebook.com")
    time.sleep(0.5)
    py.press("enter")
    time.sleep(0.5)

    try:# Checar se já esta logado
        if py.locateOnScreen("images/perfil_facebook.png", confidence=0.7):
            print("Já está logado")
            time.sleep(0.5)
        else:
            print("Não está logado")
            time.sleep(0.5)
    except Exception as e:
        print(f"Erro ao checar se está logado: {str(e)}")
        time.sleep(0.5)

    # limpar campo antes de digitar
    py.hotkey("ctrl", "a")
    time.sleep(0.5)
    py.press("backspace")
    time.sleep(0.5)

    #login facebook
    py.write(os.getenv("EMAILFACE"))
    time.sleep(0.5)
    py.press("enter")
    time.sleep(0.5)

    # Aguardar carregamento da página e tentar encontrar a barra de buscas
    print("Aguardando carregamento da página...")
    for _ in range(10):  # Tentar por até 20 segundos
        try:
            # Usar confiança mais baixa para melhorar a detecção
            pos = py.locateOnScreen("images/barra_pesquisar.png", confidence=0.7)
            if pos:
                print("Barra de buscas encontrada!")
                py.click(pos)
                time.sleep(0.5)
                print("Buscando por Python...")
                py.write("python")
                time.sleep(0.5)
                py.press("enter")
                break
        except Exception as e:
            print(f"Tentativa falhou: {str(e)}")
            time.sleep(0.5)
    else:
        print("Não foi possível encontrar a barra de buscas após várias tentativas")

    # Clicar em posts
    print("Procurando pelo botão Posts...")
    py.click(x=42, y=334)
    time.sleep(0.5)

# Executar a automação
if __name__ == "__main__":
    print("Iniciando automação...")
    
    # Primeiro fazer login
    login_facebook()
    
    #tira print dos posts
    print("\nTirando print dos posts...")
    printar_post()
    time.sleep(0.5)

    # Depois coletar posts
    print("\nIniciando coleta de posts...")
    posts = coletar_post()
    if posts:
        print("\nColeta de posts concluída com sucesso!")
    else:
        print("\nNenhum post foi coletado")